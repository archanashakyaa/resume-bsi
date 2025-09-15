from flask import Flask, render_template, request, send_file, jsonify
import subprocess
from docx import Document
import os

app = Flask(__name__)

# Configurable Ollama model
OLLAMA_MODEL = "gemma3:270m"

# Resume Enhancement Prompts (CAG-style)
resume_prompts = {
#     "summary": (
#         "You are an expert resume consultant. "
#         "Enhance the Professional Summary/Objective for a user. "
#         "Keep it polished, concise, employer-focused, and aligned with the user's actual skills and experience. "
#         "Do NOT add any experience the user does not have. "
#         "Return only one improved version of the text with no prefixes, headings, or commentary."
#     ),
#     "experience": (
#         "Enhance the Work Experience section to highlight achievements, responsibilities, and measurable impact. "
#         "Use action verbs and quantify results where possible. "
#         "Do NOT add any experience the user does not have. "
#         "Return only one improved version with no prefixes, headings, or commentary."
#     ),
#     "skills": (
#         "Improve the Skills section to be concise, organized, and impressive. "
#         "Group skills logically, include technical and soft skills, and remove redundancy. "
#         "Do NOT add any skills the user does not have. "
#         "Return one comma-separated list with no prefixes or headings."
#     ),
#     "education": (
#         "Rewrite the Education section to clearly present degrees, certifications, and relevant coursework. "
#         "Focus on what supports the user's career goals. "
#         "Do NOT add any degrees or certifications the user does not have. "
#         "Return only one improved version with no prefixes or headings."
#     ),
#     "projects": (
#         "Enhance the Projects section to present scope, technologies, contributions, and measurable impact. "
#         "Do NOT add any projects the user has not completed. "
#         "Return only one polished version with no prefixes or headings."
#     ),
#     "certifications": (
#         "Improve the Certifications section to highlight relevant certifications and their impact. "
#         "Do NOT add any certifications the user does not have. "
#         "Return only one improved version with no prefixes or headings."
#     ),
#     "achievements": (
#         "Enhance the Achievements section to highlight awards, recognitions, or accomplishments. "
#         "Do NOT add any achievements the user does not have. "
#         "Return one concise, quantifiable, professional version with no prefixes or headings."
#     ),
#     "hobbies": (
#         "Improve the Hobbies/Interests section to be professional, relevant, and reflective of skills. "
#         "Do NOT add hobbies the user does not have. "
#         "Return only one improved version with no prefixes or headings."
#     ),
    "summary": (
    "You are an expert resume consultant. Rewrite the Professional Summary/Objective section of a resume. "
    "Make it 2–3 sentences, maximum 50 words. Keep it polished, concise, and employer-focused. "
    "Only rephrase the given text — do not add new experience or skills. "
    "Output only the improved Professional Summary, no headings or commentary."
),

"experience": (
    "You are an expert resume consultant. Rewrite the Work Experience section of a resume. "
    "Highlight achievements, responsibilities, and measurable impact. Use action verbs and quantify results if mentioned. "
    "Keep each job entry concise, 2–4 lines maximum. Do not invent or add new roles. "
    "Output only the improved Work Experience, no headings or extra text."
),

"skills": (
    "You are an expert resume consultant. Rewrite the Skills section of a resume. "
    "Make it a clean, comma-separated list. Group similar skills, remove duplicates, and include only what is provided. "
    "Do not invent or add new skills. Limit to maximum 8 items. "
    "Output only the improved Skills list, nothing else."
),

"education": (
    "You are an expert resume consultant. Rewrite the Education section of a resume. "
    "List degrees, certifications, or coursework clearly and professionally. Keep each entry on one line. "
    "Do not add any new degrees or certifications. "
    "Output only the improved Education section, no headings or commentary."
),

"projects": (
    "You are an expert resume consultant. Rewrite the Projects section of a resume. "
    "Highlight scope, technologies used, contributions, and measurable results if mentioned. "
    "Do not add new projects. Keep each project 2–3 lines maximum. "
    "Output only the improved Projects section, no headings or commentary."
),

"certifications": (
    "You are an expert resume consultant. Rewrite the Certifications section of a resume. "
    "Make it concise and professional. Keep each entry one line. "
    "Do not add any new certifications. "
    "Output only the improved Certifications section, no headings or commentary."
),

"achievements": (
    "You are an expert resume consultant. Rewrite the Achievements section of a resume. "
    "Highlight awards, recognitions, or accomplishments in a concise, professional way. "
    "Do not invent achievements. Limit to maximum 3 lines. "
    "Output only the improved Achievements section, no headings or commentary."
),

"hobbies": (
    "You are an expert resume consultant. Rewrite the Hobbies/Interests section of a resume. "
    "Make it professional and relevant while keeping it brief. "
    "Only rephrase what is given — do not add hobbies. "
    "Output only the improved Hobbies section, no headings or commentary."
),
  }

# Helper Functions
def enhance_section(section_name, user_input):
    """Enhance a resume section using Ollama (CAG approach)"""
    if not user_input.strip():
        return user_input

    section_key = section_name.lower()
    if section_key not in resume_prompts:
        return user_input

    combined_prompt = f"{resume_prompts[section_key]}\n\nUser Input:\n{user_input}\n\nImproved Content:"

    try:
        result = subprocess.run(
            ["ollama", "run", OLLAMA_MODEL, combined_prompt],
            capture_output=True,
            text=True,
            check=True,
            timeout=60
        )
        enhanced_text = result.stdout.strip()

        # Remove any common prefixes or headings
        prefixes = [
            f"{section_name.title()}:",
            f"{section_name.upper()}:",
            "Summary:",
            "Improved Content:",
            "Enhanced:",
            "Here is the improved text:",
            "Here's the enhanced version:",
            "Here's a polished and concise resume summary/objective tailored for a user:"
        ]
        for p in prefixes:
            if enhanced_text.startswith(p):
                enhanced_text = enhanced_text[len(p):].strip()

        # Remove markdown bullets or extra asterisks
        enhanced_text = enhanced_text.replace("*", "").strip()

        return enhanced_text if enhanced_text else user_input

    except subprocess.TimeoutExpired:
        print(f"⏳ Timeout enhancing {section_name}")
        return user_input
    except subprocess.CalledProcessError as e:
        print(f"❌ Ollama error for {section_name}: {e.stderr}")
        return user_input
    except Exception as e:
        print(f"⚠️ Unexpected error enhancing {section_name}: {str(e)}")
        return user_input


def save_resume_docx(enhanced_resume, filename="Enhanced_Resume.docx"):
    """Save enhanced resume to DOCX file"""
    doc = Document()
    doc.add_heading("Enhanced Resume", 0)
    for section, text in enhanced_resume.items():
        if text:
            doc.add_heading(section.title(), level=1)
            doc.add_paragraph(text)
    doc.save(filename)
    return filename

# Flask Routes
@app.route("/", methods=["GET", "POST"])
def index():
    enhanced_resume = {}
    download_file = None

    if request.method == "POST":
        for section in resume_prompts.keys():
            user_input = request.form.get(section, "")
            if user_input.strip():
                enhanced_text = enhance_section(section, user_input)
                if enhanced_text:
                    enhanced_resume[section] = enhanced_text

        if enhanced_resume:
            download_file = save_resume_docx(enhanced_resume)

    return render_template("index.html", enhanced_resume=enhanced_resume, download_file=download_file)

@app.route("/enhance", methods=["POST"])
def enhance_ajax():
    """AJAX endpoint for individual section enhancement"""
    try:
        data = request.get_json()
        section_name = data.get('section')
        content = data.get('content')

        if not section_name or not content:
            return jsonify({'success': False, 'error': 'Missing section name or content'}), 400

        enhanced_content = enhance_section(section_name, content)

        return jsonify({'success': True, 'enhanced_content': enhanced_content, 'section': section_name})

    except Exception as e:
        return jsonify({'success': False, 'error': f'Server error: {str(e)}'}), 500

@app.route("/download")
def download():
    filename = "Enhanced_Resume.docx"
    if os.path.exists(filename):
        return send_file(filename, as_attachment=True)
    return "File not found.", 404

@app.route("/health")
def health_check():
    """Check Ollama and model availability"""
    try:
        result = subprocess.run(["ollama", "list"], capture_output=True, text=True, check=True)
        model_available = OLLAMA_MODEL in result.stdout
        return jsonify({'status': 'healthy', 'ollama_running': True, 'model_available': model_available})
    except Exception as e:
        return jsonify({'status': 'unhealthy', 'ollama_running': False, 'error': str(e)}), 500

# Error Handlers
@app.errorhandler(404)
def not_found(error):
    return render_template('index.html'), 404

@app.errorhandler(500)
def internal_error(error):
    return jsonify({'error': 'Internal server error'}), 500

# Main
if __name__ == "__main__":
    print("🚀 Starting Flask app on http://localhost:5000")
    app.run(debug=True, host='0.0.0.0', port=5000)
